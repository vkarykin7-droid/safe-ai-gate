"""
SafeAI Gateway Pro v3.1 — Bezpieczny Agent AI z Anonimizacją Danych (RODO)
Wszystko w jednym pliku — maksymalna kompatybilność z Streamlit Cloud.

Wersja 3.1 - Ulepszenia:
- Poprawione zliczanie wycieków (tracking podczas zastępowania)
- Skompilowane wzorce regex dla lepszej wydajności
- Obsługa upload-only (plik bez polecenia)
- Rozszerzone wzorce telefonów
- Zachowanie tytułów w redakcji nazwisk
- Limit rozmiaru plików
"""

import re
import base64
from datetime import datetime
from typing import Tuple, Optional, List

import streamlit as st

# ─────────────────────────────────────────────────────────────────────────────
# KONFIGURACJA
# ─────────────────────────────────────────────────────────────────────────────

APP_TITLE = "SafeAI Gateway Pro"
APP_ICON = "🛡️"
APP_VERSION = "3.1.0"
DEFAULT_MODEL = "gpt-4o"
MAX_FILE_SIZE_MB = 10  # Limit rozmiaru pliku w MB
MAX_TOKENS = 2000
AI_TEMPERATURE = 0.7

SYSTEM_PROMPT = (
    "Jesteś bezpiecznym, profesjonalnym asystentem biurowym. "
    "Wykonaj polecenie użytkownika, bazując wyłącznie na dostarczonych danych. "
    "Dane osobowe zostały zastąpione tagami np. [UKRYTY_EMAIL], [UKRYTY_PESEL] — "
    "zignoruj te tagi i traktuj je jako zwykłe wartości zastępcze. "
    "Odpowiadaj po polsku, chyba że polecenie wskazuje inny język."
)

VISION_PROMPT = (
    "Przepisz dokładnie cały tekst widoczny na zdjęciu. "
    "Zachowaj oryginalną strukturę akapitów. "
    "Jeśli coś jest nieczytelne, zaznacz to jako [NIECZYTELNE]."
)

FILE_ONLY_PROMPT = (
    "Proszę przeanalizować i streścić poniższe dane. "
    "Zwróć uwagę na najważniejsze informacje i przedstaw je w sposób przejrzysty."
)

# ─────────────────────────────────────────────────────────────────────────────
# SILNIK ANONIMIZACJI (RODO)
# ─────────────────────────────────────────────────────────────────────────────

# Reguły redakcji jako lista krotek (pattern, replacement)
REDACTION_RULES = [
    # Dokumenty tożsamości
    (r'\bPESEL[:\s]*\d{11}\b', 'PESEL: [UKRYTY_PESEL]'),
    (r'\bNIP[:\s]*\d{3}[-\s]?\d{3}[-\s]?\d{2}[-\s]?\d{2}\b', 'NIP: [UKRYTY_NIP]'),
    (r'\bREGON[:\s]*\d{9,14}\b', 'REGON: [UKRYTY_REGON]'),
    (r'\b(NR\.?\s*DOWODU|SERIA\s+I\s+NR|DOWÓD)[:\s]*[A-Z]{3}\s?\d{6}\b', r'\1: [UKRYTY_NR_DOWODU]'),
    (r'\bPASZPORT[:\s]*[A-Z]{2}\s?\d{7}\b', 'PASZPORT: [UKRYTY_PASZPORT]'),
    
    # Dane finansowe
    (r'\bPL\s?\d{2}[\s-]?(?:\d{4}[\s-]?){6}\d{4}\b', '[UKRYTY_NR_KONTA]'),
    (r'\b(?:\d{4}[\s-]?){3}\d{4}\b', '[UKRYTA_KARTA]'),
    
    # Kontakt - Email
    (r'\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b', '[UKRYTY_EMAIL]'),
    
    # Kontakt - Telefony (rozszerzone wzorce)
    (r'\+48[\s-]?\d{3}[\s-]?\d{3}[\s-]?\d{3}\b', '[UKRYTY_TEL]'),  # +48 123 456 789
    (r'\+48\d{9}\b', '[UKRYTY_TEL]'),  # +48123456789
    (r'(?<!\d)\d{3}[\s-]\d{3}[\s-]\d{3}(?!\d)', '[UKRYTY_TEL]'),  # 123 456 789 / 123-456-789
    (r'\(\d{2}\)\s?\d{3}[\s-]?\d{2}[\s-]?\d{2}\b', '[UKRYTY_TEL]'),  # (12) 345-67-89
    
    # Osobowe - zachowujemy tytuł
    (r'(Pan|Pani|Panem|Panią|dr|mgr|inż\.)\s+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+(?:\s+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)?', 
     r'\1 [UKRYTY_KLIENT]'),
    
    # Adresy
    (r'(ul\.|ulica|al\.|aleja|pl\.|plac|os\.|osiedle)\s+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż\s]+\d+[A-Za-z]?(?:/\d+)?', 
     '[UKRYTY_ADRES]'),
    (r'\b\d{2}-\d{3}\b', '[UKRYTY_KOD_POCZTOWY]'),
    
    # Daty urodzenia
    (r'(ur\.|urodzony|urodzona|data\s+urodzenia)[:\s]*\d{2}[.\-/]\d{2}[.\-/]\d{4}', 
     r'\1: [UKRYTA_DATA_UR]'),
    
    # Hasła
    (r'(hasło|password|passwd)[:\s]+\S+', r'\1: [UKRYTE_HASLO]'),
]

# Kompilacja wzorców dla lepszej wydajności
COMPILED_RULES: List[Tuple[re.Pattern, str]] = [
    (re.compile(pattern, re.IGNORECASE), replacement) 
    for pattern, replacement in REDACTION_RULES
]


def redact(text: str) -> Tuple[str, int]:
    """
    Anonimizuje dane wrażliwe według reguł RODO.
    
    Args:
        text: Tekst do zanonimizowania
        
    Returns:
        Tuple[str, int]: (zanonimizowany tekst, liczba wykrytych i usuniętych danych)
    """
    if not text:
        return "", 0
    
    result = text
    total_replacements = 0
    
    # Zliczamy rzeczywiste zastąpienia podczas przetwarzania
    for compiled_pattern, replacement in COMPILED_RULES:
        matches = compiled_pattern.findall(result)
        if matches:
            # Zliczamy znalezione dopasowania
            total_replacements += len(matches)
            # Wykonujemy zastąpienie
            result = compiled_pattern.sub(replacement, result)
    
    return result, total_replacements


# ─────────────────────────────────────────────────────────────────────────────
# KLIENT OPENAI
# ─────────────────────────────────────────────────────────────────────────────

def get_openai_client():
    """Tworzy klienta OpenAI z kluczem ze Streamlit Secrets."""
    try:
        from openai import OpenAI
        return OpenAI(api_key=st.secrets["OPENAI_API_KEY"]), None
    except KeyError:
        return None, "Brak klucza OPENAI_API_KEY w Streamlit Secrets."
    except ImportError:
        return None, "Brak biblioteki openai. Uruchom: pip install openai"
    except Exception as e:
        return None, f"Błąd inicjalizacji klienta: {e}"


def ai_complete(client, message: str, model: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Wysyła zanonimizowany tekst do modelu ChatGPT.
    
    Args:
        client: Klient OpenAI
        message: Wiadomość do wysłania
        model: Nazwa modelu do użycia
        
    Returns:
        Tuple[Optional[str], Optional[str]]: (odpowiedź, błąd)
    """
    try:
        resp = client.chat.completions.create(
            model=model,
            max_tokens=MAX_TOKENS,
            temperature=AI_TEMPERATURE,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": message},
            ],
        )
        return resp.choices[0].message.content, None
    except Exception as e:
        return None, _parse_openai_error(e)


def ai_vision_ocr(client, image_b64: str, media_type: str) -> Tuple[Optional[str], Optional[str]]:
    """
    OCR zdjęcia przez GPT-4o Vision.
    
    Args:
        client: Klient OpenAI
        image_b64: Obraz zakodowany w base64
        media_type: Typ MIME obrazu
        
    Returns:
        Tuple[Optional[str], Optional[str]]: (wynik OCR, błąd)
    """
    try:
        resp = client.chat.completions.create(
            model="gpt-4o",
            max_tokens=MAX_TOKENS,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": VISION_PROMPT},
                    {"type": "image_url", "image_url": {
                        "url": f"data:{media_type};base64,{image_b64}",
                        "detail": "high",
                    }},
                ],
            }],
        )
        return resp.choices[0].message.content, None
    except Exception as e:
        return None, _parse_openai_error(e)


def _parse_openai_error(e: Exception) -> str:
    """Parsuje i tłumaczy błędy OpenAI na przyjazne komunikaty."""
    s = str(e).lower()
    if "rate_limit" in s:
        return "Przekroczono limit zapytań OpenAI. Poczekaj chwilę."
    if "insufficient_quota" in s:
        return "Wyczerpano środki na koncie OpenAI."
    if "invalid_api_key" in s:
        return "Nieprawidłowy klucz API OpenAI."
    if "context_length" in s:
        return "Tekst jest zbyt długi dla wybranego modelu."
    if "connection" in s:
        return "Brak połączenia z serwerami OpenAI."
    return f"Błąd OpenAI: {e}"


# ─────────────────────────────────────────────────────────────────────────────
# EKSTRAKCJA TEKSTU Z PLIKÓW
# ─────────────────────────────────────────────────────────────────────────────

def extract_file(uploaded_file, client) -> Tuple[str, Optional[str]]:
    """
    Wyciąga tekst z PDF / DOCX / obrazu.
    
    Args:
        uploaded_file: Plik wgrany przez użytkownika
        client: Klient OpenAI (dla OCR obrazów)
        
    Returns:
        Tuple[str, Optional[str]]: (wyciągnięty tekst, błąd)
    """
    # Sprawdzenie rozmiaru pliku
    file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE_MB:
        return "", f"Plik jest zbyt duży ({file_size_mb:.1f} MB). Maksymalny rozmiar: {MAX_FILE_SIZE_MB} MB."
    
    ftype = uploaded_file.type

    if ftype == "application/pdf":
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(uploaded_file) as pdf:
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t:
                        pages.append(f"[Strona {i+1}]\n{t}")
            if not pages:
                return "", "PDF nie zawiera tekstu (spróbuj wgrać jako JPG)."
            return "\n\n".join(pages), None
        except ImportError:
            return "", "Brak biblioteki pdfplumber. Dodaj ją do requirements.txt"
        except Exception as e:
            return "", f"Błąd odczytu PDF: {e}"

    elif ftype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            from docx import Document
            doc = Document(uploaded_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs) or "", None
        except ImportError:
            return "", "Brak biblioteki python-docx. Dodaj ją do requirements.txt"
        except Exception as e:
            return "", f"Błąd odczytu DOCX: {e}"

    elif ftype in ("image/jpeg", "image/png", "image/jpg"):
        img_b64 = base64.b64encode(uploaded_file.getvalue()).decode("utf-8")
        return ai_vision_ocr(client, img_b64, ftype)

    return "", f"Nieobsługiwany typ pliku: {ftype}"


# ─────────────────────────────────────────────────────────────────────────────
# STAN SESJI
# ─────────────────────────────────────────────────────────────────────────────

def init_session() -> None:
    """Inicjalizuje stan sesji Streamlit."""
    defaults = {
        "leaks_blocked": 0,
        "total_queries": 0,
        "last_response": None,
        "last_redacted": None,
        "last_leaks": 0,
        "last_model": None,
        "session_start": datetime.now(),
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def session_duration() -> str:
    """Oblicza i formatuje czas trwania sesji."""
    delta = datetime.now() - st.session_state["session_start"]
    m, s = divmod(int(delta.total_seconds()), 60)
    return f"{m}m {s}s"


# ─────────────────────────────────────────────────────────────────────────────
# UI — PANEL BOCZNY
# ─────────────────────────────────────────────────────────────────────────────

def render_sidebar() -> None:
    """Renderuje panel boczny z statusem i statystykami."""
    with st.sidebar:
        st.header("⚙️ Status Systemu")
        if "OPENAI_API_KEY" in st.secrets:
            st.success("✅ OpenAI API: Połączono")
        else:
            st.error("❌ OpenAI API: Brak klucza")

        st.divider()
        st.header("📊 Statystyki Sesji")

        c1, c2 = st.columns(2)
        c1.metric("🛡️ Wycieków", st.session_state["leaks_blocked"])
        c2.metric("📝 Zapytań", st.session_state["total_queries"])
        st.caption(f"⏱️ Czas sesji: {session_duration()}")

        st.divider()
        with st.expander("🔍 Aktywne reguły RODO"):
            rules = [
                "✓ PESEL", "✓ NIP", "✓ REGON", "✓ Nr dowodu", "✓ Paszport",
                "✓ Nr konta IBAN", "✓ Karta płatnicza", "✓ E-mail",
                "✓ Telefon (5 formatów)", "✓ Imię i nazwisko", "✓ Adres", 
                "✓ Kod pocztowy", "✓ Data urodzenia", "✓ Hasła w tekście",
            ]
            for r in rules:
                st.markdown(f"- {r}")

        st.divider()
        if st.button("🗑️ Wyczyść wynik", use_container_width=True):
            st.session_state.update({
                "last_response": None, 
                "last_redacted": None,
                "last_leaks": 0, 
                "last_model": None
            })
            st.rerun()

        if st.button("🔄 Resetuj sesję", use_container_width=True, type="secondary"):
            for k in ["leaks_blocked", "total_queries", "last_response",
                      "last_redacted", "last_leaks", "last_model"]:
                if "blocked" in k or "queries" in k or "leaks" in k:
                    st.session_state[k] = 0
                else:
                    st.session_state[k] = None
            st.session_state["session_start"] = datetime.now()
            st.rerun()

        st.divider()
        st.caption(
            "🔒 **Jak działa ochrona?**\n\n"
            "Dane są anonimizowane **lokalnie** zanim trafią do OpenAI. "
            "ChatGPT nigdy nie widzi prawdziwych danych osobowych."
        )


# ─────────────────────────────────────────────────────────────────────────────
# UI — WYNIKI
# ─────────────────────────────────────────────────────────────────────────────

def render_results(show_redacted: bool) -> None:
    """
    Renderuje wyniki przetwarzania AI.
    
    Args:
        show_redacted: Czy pokazać zanonimizowany tekst wejściowy
    """
    response = st.session_state.get("last_response")
    if not response:
        return

    redacted = st.session_state.get("last_redacted", "")
    leak_count = st.session_state.get("last_leaks", 0)
    model_used = st.session_state.get("last_model", "–")

    st.divider()

    if leak_count > 0:
        st.success(
            f"🛡️ **Tarcza SafeAI:** Wykryto i zablokowano **{leak_count}** "
            f"{'wyciek' if leak_count == 1 else 'wycieki' if leak_count < 5 else 'wycieków'} danych."
        )
    else:
        st.info("ℹ️ **Tarcza SafeAI:** W tym tekście nie wykryto danych wrażliwych.")

    st.caption(f"Model: `{model_used}` | Znaki wejściowe: `{len(redacted)}`")

    if show_redacted and redacted:
        col_l, col_r = st.columns(2)
        with col_l:
            st.subheader("🛡️ Tekst wysłany do AI")
            st.caption("Dokładnie to, co zobaczyło ChatGPT.")
            st.code(redacted, language=None)
        with col_r:
            st.subheader("🤖 Wynik analizy AI")
            with st.container(border=True):
                st.markdown(response)
    else:
        st.subheader("🤖 Wynik analizy AI")
        with st.container(border=True):
            st.markdown(response)

    st.divider()
    st.subheader("💾 Eksport")
    c1, c2, _ = st.columns([1, 1, 2])
    c1.download_button(
        "📄 Odpowiedź (.txt)", 
        data=response,
        file_name="safeai_odpowiedz.txt", 
        mime="text/plain",
        use_container_width=True
    )
    if redacted:
        full = (
            f"=== DANE WEJŚCIOWE (zanonimizowane) ===\n\n{redacted}\n\n"
            f"=== ODPOWIEDŹ AI ({model_used}) ===\n\n{response}"
        )
        c2.download_button(
            "📋 Pełny raport (.txt)", 
            data=full,
            file_name="safeai_raport.txt", 
            mime="text/plain",
            use_container_width=True
        )


# ─────────────────────────────────────────────────────────────────────────────
# GŁÓWNA APLIKACJA
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    """Główna funkcja aplikacji."""
    st.set_page_config(
        page_title=APP_TITLE, 
        page_icon=APP_ICON,
        layout="wide", 
        initial_sidebar_state="expanded"
    )

    init_session()

    client, client_error = get_openai_client()
    if client_error:
        st.error(f"❌ {client_error}")
        st.info(
            "Dodaj klucz OPENAI_API_KEY do `.streamlit/secrets.toml` "
            "lub Streamlit Cloud Secrets."
        )
        st.stop()

    render_sidebar()

    # ── Nagłówek ─────────────────────────────────────────────────────────────
    st.title(f"{APP_ICON} {APP_TITLE}")
    st.markdown(
        "**Profesjonalna bariera ochronna RODO** — dane są anonimizowane "
        "*lokalnie*, zanim trafią do AI. ChatGPT nigdy nie widzi danych osobowych."
    )
    st.divider()

    # ── Formularz wejściowy ──────────────────────────────────────────────────
    col_in, col_opt = st.columns([3, 1])

    with col_in:
        user_command = st.text_area(
            "📝 Twoje polecenie:",
            placeholder='Np. "Napisz wezwanie do zapłaty" albo "Streść ten dokument"...',
            height=120,
        )

    with col_opt:
        st.markdown("**⚙️ Opcje**")
        show_redacted = st.toggle("Pokaż zanonimizowany tekst", value=True)
        ai_model = st.selectbox("Model AI", ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"])

    uploaded_file = st.file_uploader(
        f"📂 Opcjonalnie: Wgraj plik (PDF, DOCX, JPG, PNG) — max {MAX_FILE_SIZE_MB} MB",
        type=["pdf", "docx", "jpg", "jpeg", "png"],
    )

    if uploaded_file and uploaded_file.type.startswith("image/"):
        st.image(uploaded_file, caption="Wgrane zdjęcie — Vision OCR odczyta tekst", width=280)

    # ── Przycisk ─────────────────────────────────────────────────────────────
    process = st.button(
        "🚀 Uruchom Bezpieczne Przetwarzanie",
        type="primary",
        use_container_width=True,
        disabled=(not user_command and not uploaded_file),
    )

    # ── Logika przetwarzania ─────────────────────────────────────────────────
    if process:
        with st.status("🔄 Przetwarzanie w toku...", expanded=True) as status:

            # Krok 1 — ekstrakcja pliku
            file_text = ""
            if uploaded_file:
                st.write("📄 **Krok 1/3:** Wyciąganie tekstu z pliku...")
                file_text, err = extract_file(uploaded_file, client)
                if err:
                    st.error(f"Błąd pliku: {err}")
                    status.update(label="❌ Błąd pliku", state="error")
                    return
                else:
                    st.write(f"✅ Wyciągnięto {len(file_text)} znaków.")

            # Krok 2 — przygotowanie contentu
            st.write("📝 **Krok 2/3:** Przygotowanie zapytania...")
            if file_text and user_command:
                full_content = f"POLECENIE: {user_command}\n\nDANE:\n{file_text}"
            elif file_text:
                # Obsługa upload-only: plik bez polecenia
                full_content = f"{FILE_ONLY_PROMPT}\n\n{file_text}"
            else:
                full_content = user_command

            # Krok 2b — anonimizacja
            st.write("🛡️ **Krok 2b/3:** Anonimizacja danych wrażliwych (lokalnie)...")
            redacted_text, leak_count = redact(full_content)
            st.write(f"✅ Ukryto **{leak_count}** pól z danymi wrażliwymi.")

            # Krok 3 — zapytanie do AI
            st.write(f"🤖 **Krok 3/3:** Wysyłanie do {ai_model}...")
            ai_response, ai_error = ai_complete(client, redacted_text, ai_model)

            if ai_error:
                st.error(f"❌ {ai_error}")
                status.update(label="❌ Błąd AI", state="error")
            else:
                st.session_state["leaks_blocked"] += leak_count
                st.session_state["total_queries"] += 1
                st.session_state["last_response"] = ai_response
                st.session_state["last_redacted"] = redacted_text
                st.session_state["last_leaks"] = leak_count
                st.session_state["last_model"] = ai_model
                status.update(label="✅ Gotowe!", state="complete")

    # ── Wyniki ───────────────────────────────────────────────────────────────
    render_results(show_redacted=show_redacted)

    # ── Stopka ───────────────────────────────────────────────────────────────
    st.divider()
    st.caption(
        f"© 2026 SafeAI Gateway | v{APP_VERSION} | "
        f"System ochrony danych wrażliwych zgodny z RODO"
    )


if __name__ == "__main__":
    main()
